

import os
import uuid
import shutil
from typing import List, Dict, Any, Optional, Union
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv


import numpy as np
from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks, Depends, Query, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import APIKeyHeader
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session
from docx import Document
import nltk
from nltk.tokenize import sent_tokenize
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer

#auth
from auth import api_key_header, validate_api_key

# Import database models (to be implemented in models.py)
from models import get_db, engine
import models

# Download required NLTK resources
nltk.download('punkt')
nltk.download('stopwords')

# Initialize the FastAPI app
app = FastAPI(
    title="Journal Plagiarism Checker API",
    description="API for detecting plagiarism in academic journal submissions",
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Update with specific origins in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# API key authentication
API_KEY_NAME = "X-API-Key"
api_key_header = APIKeyHeader(name=API_KEY_NAME)

# Near the top of your file, after imports
from database import Base, engine

# Initialize database tables
@app.on_event("startup")
async def startup():
    Base.metadata.create_all(bind=engine)
    print("Database tables created")
# Initialize models 

models.Base.metadata.create_all(bind=engine)

# Initialize the sentence transformer model with error handling
try:
    print("Loading sentence transformer model...")
    model = SentenceTransformer('paraphrase-MiniLM-L6-v2')
    print("Model loaded successfully")
except Exception as e:
    print(f"Error loading sentence transformer model: {str(e)}")
    # Fallback to a simpler model if available
    try:
        print("Attempting to load fallback model...")
        model = SentenceTransformer('distilbert-base-nli-mean-tokens')
        print("Fallback model loaded successfully")
    except Exception as e2:
        print(f"Error loading fallback model: {str(e2)}")
        print("WARNING: Using dummy embedding function instead of transformer model")
        # Define a dummy embedding function that returns zeros
        def create_embeddings(text_chunks):
            print("Using dummy embeddings")
            return [[0.0] * 10 for _ in text_chunks]
# Create upload directory if it doesn't exist
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# Models for API request/response
class DocumentResponse(BaseModel):
    id: str
    filename: str
    upload_date: datetime
    size: int
    status: str
    
    class Config:
        orm_mode = True

class PlagiarismCheckRequest(BaseModel):
    document_id: str
    threshold: float = Field(0.7, description="Similarity threshold (0.0 to 1.0)")

class PlagiarismMatch(BaseModel):
    document_id: str
    document_title: str
    similarity_score: float
    matched_segments: List[Dict[str, Any]]
    
    class Config:
        orm_mode = True

class PlagiarismResult(BaseModel):
    id: str
    document_id: str
    overall_similarity: float
    matches: List[PlagiarismMatch]
    created_at: datetime
    
    class Config:
        orm_mode = True

# Helper functions
def extract_text_from_docx(file_path: str) -> str:
    """Extract text content from a DOCX file."""
    doc = Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

def preprocess_text(text: str) -> List[str]:
    """Preprocess text for plagiarism detection."""
    # Split text into sentences
    sentences = sent_tokenize(text)
    
    # Remove stopwords
    stop_words = set(stopwords.words('english'))
    processed_sentences = []
    
    for sentence in sentences:
        # Lowercase and remove stopwords
        words = [word.lower() for word in sentence.split() if word.lower() not in stop_words]
        if words:
            processed_sentences.append(" ".join(words))
    
    return processed_sentences

def create_embeddings(text_chunks: List[str]):
    """Create embeddings for text chunks using sentence transformer."""
    return model.encode(text_chunks)

def check_plagiarism_background(document_id: str, threshold: float, db: Session):
    """Background task to check document for plagiarism."""
    document = None
    try:
        # Get the document from the database
        document = db.query(models.Document).filter(models.Document.id == document_id).first()
        if not document:
            print(f"Document {document_id} not found")
            return
        
        print(f"Starting processing for document: {document.filename}")
        
        # Update document status
        document.status = "processing"
        db.commit()
        
        # Get the document file path
        file_path = UPLOAD_DIR / document.file_path
        print(f"Reading file from: {file_path}")
        
        # Extract and preprocess text
        try:
            text = extract_text_from_docx(str(file_path))
            print(f"Extracted {len(text)} characters of text")
        except Exception as e:
            print(f"Error extracting text: {str(e)}")
            raise
            
        try:
            processed_text = preprocess_text(text)
            print(f"Processed into {len(processed_text)} text segments")
            
            if not processed_text:
                print("Warning: No text segments after preprocessing")
                # Create a minimal processed text to avoid failures
                processed_text = ["Empty document"]
        except Exception as e:
            print(f"Error preprocessing text: {str(e)}")
            raise
        
        # Create embeddings
        try:
            document_embeddings = create_embeddings(processed_text)
            print(f"Created {len(document_embeddings)} embeddings")
        except Exception as e:
            print(f"Error creating embeddings: {str(e)}")
            raise
        
        # Get all other documents
        try:
            other_documents = db.query(models.Document).filter(
                models.Document.id != document_id,
                models.Document.status == "processed"
            ).all()
            print(f"Found {len(other_documents)} other documents to compare against")
        except Exception as e:
            print(f"Error querying other documents: {str(e)}")
            raise
        
        # Create a new plagiarism check result
        result = models.PlagiarismResult(
            id=str(uuid.uuid4()),
            document_id=document_id,
            overall_similarity=0.0,
            created_at=datetime.now()
        )
        db.add(result)
        db.flush()
        print(f"Created new result record with ID: {result.id}")
        
        overall_similarities = []
        
        # Check against each document
        for other_doc in other_documents:
            try:
                print(f"Comparing against document: {other_doc.filename} (ID: {other_doc.id})")
                other_file_path = UPLOAD_DIR / other_doc.file_path
                
                # Skip if file doesn't exist
                if not os.path.exists(other_file_path):
                    print(f"Warning: File not found for document {other_doc.id}: {other_file_path}")
                    continue
                
                other_text = extract_text_from_docx(str(other_file_path))
                other_processed_text = preprocess_text(other_text)
                
                # Skip if no processed text
                if not other_processed_text:
                    print(f"Warning: No processed text for document {other_doc.id}")
                    continue
                
                other_embeddings = create_embeddings(other_processed_text)
                
                # Calculate cosine similarity between document chunks
                max_similarities = []
                matched_segments = []
                
                for i, emb1 in enumerate(document_embeddings):
                    chunk_similarities = []
                    for j, emb2 in enumerate(other_embeddings):
                        # Calculate cosine similarity
                        similarity = float(cosine_similarity([emb1], [emb2])[0][0])
                        if similarity >= threshold:
                            chunk_similarities.append({
                                "similarity": similarity,
                                "source_segment": processed_text[i],
                                "matched_segment": other_processed_text[j],
                                "source_index": i,
                                "matched_index": j
                            })
                    
                    # Get highest similarity for this chunk
                    if chunk_similarities:
                        max_similarity = max(chunk_similarities, key=lambda x: x["similarity"])
                        max_similarities.append(max_similarity["similarity"])
                        matched_segments.append(max_similarity)
                
                # Calculate overall similarity for this document
                if max_similarities:
                    doc_similarity = sum(max_similarities) / len(max_similarities)
                    overall_similarities.append(doc_similarity)
                    print(f"Similarity with {other_doc.filename}: {doc_similarity:.4f}")
                    
                    # Add match if above threshold
                    if doc_similarity >= threshold:
                        match = models.PlagiarismMatch(
                            id=str(uuid.uuid4()),
                            result_id=result.id,
                            matched_document_id=other_doc.id,
                            similarity_score=float(doc_similarity)
                        )
                        db.add(match)
                        db.flush()
                        print(f"Added match record with ID: {match.id}")
                        
                        # Add matched segments
                        for segment in matched_segments:
                            match_segment = models.MatchedSegment(
                                id=str(uuid.uuid4()),
                                match_id=match.id,
                                source_segment=segment["source_segment"],
                                matched_segment=segment["matched_segment"],
                                similarity=segment["similarity"],
                                source_index=segment["source_index"],
                                matched_index=segment["matched_index"]
                            )
                            db.add(match_segment)
                
            except Exception as e:
                print(f"Error comparing with document {other_doc.id}: {str(e)}")
                # Continue with next document rather than failing entire process
                continue
        
        # Calculate overall similarity across all documents
        if overall_similarities:
            result.overall_similarity = float(max(overall_similarities))
            print(f"Overall similarity: {result.overall_similarity:.4f}")
        else:
            print("No similarities found with any document")
        
        # Update document status
        document.status = "processed"
        db.commit()
        print(f"Document {document.filename} processed successfully")
        
    except Exception as e:
        print(f"Error in plagiarism check: {str(e)}")
        import traceback
        traceback.print_exc()
        
        # Update document status to error
        try:
            if document:
                document.status = "error"
                db.commit()
                print(f"Updated document {document_id} status to 'error'")
        except Exception as inner_e:
            print(f"Error updating document status: {str(inner_e)}")

# API Endpoints

@app.get("/")
async def root():
    return {"message": "Journal Plagiarism Checker API"}

@app.post("/documents/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    api_key: str = Depends(api_key_header)
):
    # Validate API key
    validate_api_key(api_key)
    
    # Check file type
    if not file.filename.endswith('.docx'):
        raise HTTPException(
            status_code=400,
            detail="Only .docx files are supported"
        )
    
    # Generate unique ID and file path
    document_id = str(uuid.uuid4())
    file_extension = os.path.splitext(file.filename)[1]
    file_path = f"{document_id}{file_extension}"
    full_path = UPLOAD_DIR / file_path
    
    # Save the file
    with open(full_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # Create document record
    document = models.Document(
        id=document_id,
        filename=file.filename,
        file_path=file_path,
        upload_date=datetime.now(),
        size=os.path.getsize(full_path),
        status="uploaded"
    )
    
    db.add(document)
    db.commit()
    db.refresh(document)
    
    return document

@app.post("/documents/check", status_code=status.HTTP_202_ACCEPTED)
async def check_plagiarism(
    request: PlagiarismCheckRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    api_key: str = Depends(api_key_header)
):
    # Validate API key
    # validate_api_key(api_key, db)
    
    # Check if document exists
    document = db.query(models.Document).filter(models.Document.id == request.document_id).first()
    if not document:
        raise HTTPException(
            status_code=404,
            detail="Document not found"
        )
    
    # Start plagiarism check in background
    background_tasks.add_task(
        check_plagiarism_background,
        document_id=request.document_id,
        threshold=request.threshold,
        db=db
    )
    
    return {
        "message": "Plagiarism check started",
        "document_id": request.document_id
    }

@app.get("/documents/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: str,
    db: Session = Depends(get_db),
    api_key: str = Depends(api_key_header)
):
    # Validate API key
    # validate_api_key(api_key, db)
    
    document = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not document:
        raise HTTPException(
            status_code=404,
            detail="Document not found"
        )
    
    return document

@app.get("/results/{document_id}", response_model=Optional[PlagiarismResult])
async def get_results(
    document_id: str,
    db: Session = Depends(get_db),
    api_key: str = Depends(api_key_header)
):
    # Validate API key
    # validate_api_key(api_key, db)
    
    # Check if document exists
    document = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not document:
        raise HTTPException(
            status_code=404,
            detail="Document not found"
        )
    
    # Get the latest plagiarism result
    result = db.query(models.PlagiarismResult).filter(
        models.PlagiarismResult.document_id == document_id
    ).order_by(models.PlagiarismResult.created_at.desc()).first()
    
    if not result:
        return None
    
    # Get matches
    matches = db.query(models.PlagiarismMatch).filter(
        models.PlagiarismMatch.result_id == result.id
    ).all()
    
    # Construct response
    response_matches = []
    for match in matches:
        # Get matched document
        matched_doc = db.query(models.Document).filter(
            models.Document.id == match.matched_document_id
        ).first()
        
        # Get segments
        segments = db.query(models.MatchedSegment).filter(
            models.MatchedSegment.match_id == match.id
        ).all()
        
        segments_data = [
            {
                "source_segment": segment.source_segment,
                "matched_segment": segment.matched_segment,
                "similarity": segment.similarity
            }
            for segment in segments
        ]
        
        response_matches.append(
            PlagiarismMatch(
                document_id=matched_doc.id,
                document_title=matched_doc.filename,
                similarity_score=match.similarity_score,
                matched_segments=segments_data
            )
        )
    
    return PlagiarismResult(
        id=result.id,
        document_id=document_id,
        overall_similarity=result.overall_similarity,
        matches=response_matches,
        created_at=result.created_at
    )

@app.get("/documents", response_model=List[DocumentResponse])
async def list_documents(
    skip: int = Query(0, ge=0),
    limit: int = Query(100, ge=1, le=100),
    db: Session = Depends(get_db),
    api_key: str = Depends(api_key_header)
):
    # Validate API key
    # validate_api_key(api_key, db)
    
    documents = db.query(models.Document).offset(skip).limit(limit).all()
    return documents

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)