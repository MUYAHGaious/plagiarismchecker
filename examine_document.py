import sys
import sqlite3
from docx import Document
import os

def examine_document(document_id):
    """Examine a document to see what might be causing issues."""
    conn = sqlite3.connect('plagiarism_checker.db')
    cursor = conn.cursor()
    
    # Get document info
    cursor.execute("SELECT id, filename, file_path, status FROM documents WHERE id = ?", 
                  (document_id,))
    doc = cursor.fetchone()
    
    if not doc:
        print(f"Document {document_id} not found")
        conn.close()
        return
    
    doc_id, filename, file_path, status = doc
    
    print(f"Document ID: {doc_id}")
    print(f"Filename: {filename}")
    print(f"Status: {status}")
    
    # Check if file exists
    full_path = os.path.join("uploads", file_path)
    if not os.path.exists(full_path):
        print(f"Error: File does not exist at {full_path}")
        conn.close()
        return
    
    print(f"File exists at {full_path}, size: {os.path.getsize(full_path)} bytes")
    
    # Try to open and read the document
    try:
        doc = Document(full_path)
        paragraph_count = len(doc.paragraphs)
        
        print(f"Document has {paragraph_count} paragraphs")
        
        # Print sample of text
        text = "\n".join([p.text for p in doc.paragraphs])
        print(f"Total text length: {len(text)} characters")
        if text:
            print("Sample text (first 200 chars):")
            print(text[:200])
        else:
            print("Warning: Document contains no text")
        
        # Check for other content
        tables_count = len(doc.tables)
        print(f"Document has {tables_count} tables")
        
        has_complex_content = False
        if tables_count > 0 or any(run.bold or run.italic or len(run._element.xpath('.//w:drawing')) > 0 
                                   for paragraph in doc.paragraphs 
                                   for run in paragraph.runs):
            has_complex_content = True
        
        print(f"Contains complex content (tables, formatting, images): {has_complex_content}")
        
    except Exception as e:
        print(f"Error examining document content: {str(e)}")
    
    # Check if there are any results for this document
    cursor.execute("SELECT id, overall_similarity FROM plagiarism_results WHERE document_id = ?", 
                  (document_id,))
    results = cursor.fetchall()
    
    if results:
        print(f"Document has {len(results)} plagiarism check results:")
        for result_id, similarity in results:
            print(f"  Result {result_id}: Similarity {similarity}")
    else:
        print("Document has no plagiarism check results")
    
    conn.close()

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Please provide a document ID to examine")
        print("Usage: python examine_document.py <document_id>")
        sys.exit(1)
    
    examine_document(sys.argv[1])